#!/usr/bin/env python3
import os
import sqlite3
from pathlib import Path
from typing import Optional, List, Tuple

from nltk.tokenize import sent_tokenize
import pytesseract
from pdf2image import convert_from_path
from PyPDF2 import PdfReader
import docx
import openpyxl

from . import logger
from .config import config


# === SQLite state DB helpers ===
def initialize_state_db():
    os.makedirs(os.path.dirname(config.STATE_DB_PATH), exist_ok=True)

    logger.info(f"Using sqlite database '{config.STATE_DB_PATH}'")

    conn = sqlite3.connect(config.STATE_DB_PATH)
    c = conn.cursor()
    c.execute(
        """
        CREATE TABLE IF NOT EXISTS files (
            path TEXT PRIMARY KEY,
            last_modified REAL
        )
    """
    )
    conn.commit()
    conn.close()


def get_stored_timestamp(relpath: Path) -> Optional[float]:
    conn = sqlite3.connect(config.STATE_DB_PATH)
    c = conn.cursor()
    c.execute("SELECT last_modified FROM files WHERE path = ?", (str(relpath),))
    row = c.fetchone()
    conn.close()
    return row[0] if row else None


def set_stored_timestamp(relpath: Path, ts: float):
    conn = sqlite3.connect(config.STATE_DB_PATH)
    c = conn.cursor()
    c.execute("REPLACE INTO files (path, last_modified) VALUES (?, ?)", (str(relpath), ts))
    conn.commit()
    conn.close()


def delete_stored_file(relpath: Path):
    conn = sqlite3.connect(config.STATE_DB_PATH)
    c = conn.cursor()
    c.execute("DELETE FROM files WHERE path = ?", (str(relpath),))
    conn.commit()
    conn.close()


def rename_folder(srcpath: Path, destpath: Path):
    conn = sqlite3.connect(config.STATE_DB_PATH)
    c = conn.cursor()
    c.execute("UPDATE files SET path=? WHERE path LIKE ?", (str(destpath), f"%{srcpath}"))
    conn.commit()
    conn.close()


def rename_stored_file(srcpath: Path, destpath: Path):
    conn = sqlite3.connect(config.STATE_DB_PATH)
    c = conn.cursor()
    c.execute("UPDATE files SET path=? WHERE path=?", (str(destpath), str(srcpath)))
    conn.commit()
    conn.close()


def list_stored_files(absolute: bool = False) -> list[Path]:
    conn = sqlite3.connect(config.STATE_DB_PATH)
    c = conn.cursor()
    c.execute("SELECT path FROM files")
    rows = c.fetchall()
    conn.close()

    files_list = []
    for (stored_path,) in rows:
        relpath = Path(stored_path)
        if absolute:
            files_list.append(config.DOCS_PATH / relpath)
        else:
            files_list.append(relpath)

    return files_list


# === Text extraction per filetype ===
def extract_text_from_pdf(path: Path) -> str:
    text_chunks = []
    len_text = 0
    reader = PdfReader(path)
    nb_pages = len(reader.pages)
    try:
        logger.info(f"Reading {nb_pages} pages pdf file")
        for page in reader.pages:
            txt = page.extract_text() or ""
            text_chunks.append(txt)
            len_text += len(txt)

        full_text = "\n".join(text_chunks).strip()

        # If nearly empty, fallback to OCR
        if len_text < 10 * nb_pages:
            logger.info("PDF text is too short; falling back to OCR")
            return ocr_pdf(path, nb_pages), {"ocr_used": True}

        return full_text, {"ocr_used": False}

    except Exception as e:
        logger.warning(f"Error extracting PDF text: {e}. Using OCR.")
        return ocr_pdf(path, nb_pages), {"ocr_used": True}


def ocr_pdf(path: Path, nb_pages: int) -> str:
    text = []
    logger.info("OCR")
    try:
        # Convert each page to an image
        for k_page in range(1, 1 + nb_pages):
            images = convert_from_path(path, first_page=k_page, last_page=k_page, dpi=300)
            for img in images:
                txt = pytesseract.image_to_string(img, lang=config.OCR_LANG)
                text.append(txt)
    except Exception as e:
        logger.error(f"OCR failed: {e}")
    return "\n".join(text).strip()


def extract_text_from_docx(path: Path) -> str:
    try:
        doc = docx.Document(str(path))
        page_count = sum(p.contains_page_break for p in doc.paragraphs) + 1
        logger.info(f"Reading {page_count} pages doc file")
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs).strip(), {"ocr_used": False}
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        return "", {"ocr_used": False}


def extract_text_from_xlsx(path: Path) -> str:
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        logger.info(f"Reading {len(wb.worksheets)} pages excel file")
        all_text = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = [str(cell) for cell in row if cell is not None]
                if row_text:
                    all_text.append(" ".join(row_text))
        return "\n".join(all_text).strip(), {"ocr_used": False}
    except Exception as e:
        logger.error(f"Error reading XLSX: {e}")
        return "", {"ocr_used": False}


def extract_text_from_md_or_txt(path: Path) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read(), {"ocr_used": False}
    except Exception as e:
        logger.error(f"Error reading text file: {e}")
        return "", {"ocr_used": False}


def extract_text(path: Path) -> Tuple[str, dict]:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext in (".docx",):
        return extract_text_from_docx(path)
    elif ext in (".xlsx", ".xlsm", ".xls"):
        return extract_text_from_xlsx(path)
    elif ext in (".md", ".txt"):
        return extract_text_from_md_or_txt(path)
    else:
        logger.warning("Unsupported extension (skipping text extraction)")
        return "", {}


def chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    """
    Splits text into overlapping chunks of ~chunk_size characters, aligned on sentences.
    """
    sentences = sent_tokenize(text)
    chunks = []
    current_chunk = ""
    logger.info("Building chunks")
    for sent in sentences:
        if len(current_chunk) + len(sent) + 1 <= chunk_size:
            current_chunk += " " + sent if current_chunk else sent
        else:
            chunks.append(current_chunk)
            # Start new chunk: include overlap
            overlap_text = (
                current_chunk[-chunk_overlap:]
                if chunk_overlap < len(current_chunk)
                else current_chunk
            )
            current_chunk = overlap_text + " " + sent

    if current_chunk:
        chunks.append(current_chunk)
    return chunks
