from typing import Iterable, Tuple

import docx
from tqdm import tqdm

from .. import logger
from .Document import Document


class DocDocument(Document):
    def iterate_raw_text(self) -> Iterable[Tuple[str, dict]]:
        try:
            doc = docx.Document(str(self.get_abs_path()))
        except Exception:
            logger.warning("Error while reading the file. Skipping")
            return None, {"ocr_used": False}

        page_count = sum(p.contains_page_break for p in doc.paragraphs) + 1
        logger.info(f"Reading {page_count} pages doc file")
        avct = -1
        for k_page, p in enumerate(tqdm(doc.paragraphs)):
            new_avct = int(k_page / page_count * 100)
            if new_avct != avct:
                logger.info(f"Lecture page {k_page+1}/{page_count}")
                avct = new_avct

            yield "\n".join(p.text).strip(), {"ocr_used": False}
